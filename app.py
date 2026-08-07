"""
Orçamentos — Papel para PDF
----------------------------
Aplicação Streamlit que permite tirar uma foto a um orçamento escrito à mão
e transformá-la automaticamente num PDF profissional, pronto a enviar ao
cliente.

Fluxo em 3 passos, pensado para ser usado a partir de um telemóvel por
pessoas com pouca experiência em aplicações digitais:
    1. Foto do orçamento em papel
    2. Confirmação e edição dos dados extraídos
    3. Descarregar o PDF final
"""

import json
import os
import re
import subprocess
import tempfile
from copy import deepcopy
from datetime import datetime

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from google import genai
from PIL import Image

load_dotenv()

# =========================================================================
# CONFIGURAÇÃO
# =========================================================================

try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except Exception:
    API_KEY = os.getenv("API_KEY")

try:
    PASSWORD_SISTEMA = st.secrets["APP_PASSWORD"]
except Exception:
    PASSWORD_SISTEMA = os.getenv("APP_PASSWORD")

TEMPLATE_PATH = "template.docx"

COLUNAS_TABELA = ["Designação", "Unidade", "Quantidade", "Preço Unitário (€)"]

NOME_MODELO_IA = "gemini-3.5-flash-lite"


# =========================================================================
# ESTILO
# =========================================================================

def aplicar_estilo():
    """Aplica o CSS que dá à app um aspeto profissional e adaptado a ecrãs
    de telemóvel e a utilizadores com menos à-vontade digital: texto maior,
    botões grandes e fáceis de tocar, e um indicador de passos claro."""
    st.markdown(
        """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}

        .stApp {
            max-width: 640px;
            margin: 0 auto;
        }

        html, body, [class*="css"] {
            font-size: 17px;
        }

        h1 {
            font-size: 1.8rem !important;
            text-align: center;
            color: #1B4965;
            margin-bottom: 0 !important;
        }
        h2 {
            font-size: 1.4rem !important;
            color: #1B4965;
            margin-bottom: 0.3rem !important;
        }
        h4 {
            color: #1B4965;
        }

        p, label, li, .stMarkdown, .stCaption {
            font-size: 1rem;
        }

        .stButton > button, .stDownloadButton > button {
            font-size: 1.05rem;
            font-weight: 600;
            padding: 0.7rem 1rem;
            border-radius: 12px;
            min-height: 3.1rem;
            border: none;
        }

        .stTextInput input, .stTextArea textarea {
            font-size: 1.05rem;
            border-radius: 10px;
        }

        div[data-testid="stMetric"] {
            background: #EEF2F5;
            padding: 1rem;
            border-radius: 14px;
        }

        /* Indicador de passos */
        .passos-container {
            display: flex;
            align-items: flex-start;
            justify-content: center;
            margin: 0.5rem 0 1.8rem 0;
        }
        .passo {
            display: flex;
            flex-direction: column;
            align-items: center;
            width: 76px;
        }
        .passo-numero {
            width: 34px;
            height: 34px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 700;
            font-size: 0.95rem;
            background: #E3E8EC;
            color: #6B7785;
        }
        .passo.ativo .passo-numero {
            background: #D9A441;
            color: #24313D;
        }
        .passo.concluido .passo-numero {
            background: #1B4965;
            color: #FFFFFF;
        }
        .passo-nome {
            font-size: 0.78rem;
            margin-top: 5px;
            color: #6B7785;
            text-align: center;
        }
        .passo.ativo .passo-nome {
            color: #1B4965;
            font-weight: 700;
        }
        .passo-linha {
            height: 3px;
            width: 34px;
            background: #E3E8EC;
            margin-top: 17px;
        }
        .passo-linha.concluido {
            background: #1B4965;
        }

        /* Cartão de acesso */
        .login-card {
            max-width: 380px;
            margin: 3rem auto 0 auto;
            padding: 2rem 1.5rem;
            border-radius: 16px;
            background: #EEF2F5;
            text-align: center;
        }

        .dica-caixa {
            background: #FBF3E2;
            border-left: 4px solid #D9A441;
            padding: 0.7rem 1rem;
            border-radius: 8px;
            font-size: 0.95rem;
            margin-bottom: 1rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def indicador_passos(passo_atual):
    """Mostra 1 → 2 → 3 no topo da página, para a pessoa saber sempre em
    que ponto do processo está."""
    nomes = ["Foto", "Confirmar", "Descarregar"]
    partes = ['<div class="passos-container">']
    for i, nome in enumerate(nomes, start=1):
        if i < passo_atual:
            estado, rotulo = "concluido", "✓"
        elif i == passo_atual:
            estado, rotulo = "ativo", str(i)
        else:
            estado, rotulo = "pendente", str(i)

        partes.append(
            f'<div class="passo {estado}">'
            f'<div class="passo-numero">{rotulo}</div>'
            f'<div class="passo-nome">{nome}</div>'
            f"</div>"
        )
        if i < len(nomes):
            linha_estado = "concluido" if i < passo_atual else ""
            partes.append(f'<div class="passo-linha {linha_estado}"></div>')
    partes.append("</div>")
    st.markdown("".join(partes), unsafe_allow_html=True)


# =========================================================================
# FUNÇÕES DE APOIO (dados e formatação)
# =========================================================================

def parse_numero(valor, default=0.0):
    """Converte um valor (número, texto com vírgula ou ponto, ou vazio)
    num número, de forma tolerante a diferentes formatos."""
    if valor is None or (isinstance(valor, float) and pd.isna(valor)):
        return default
    if isinstance(valor, (int, float)):
        return float(valor)
    texto = str(valor).strip()
    if not texto:
        return default
    texto = re.sub(r"[€\s]", "", texto).replace(",", ".")
    try:
        return float(texto)
    except ValueError:
        return default


def formatar_numero(valor):
    return f"{valor:.2f}"


def formatar_euro(valor):
    return f"{valor:.2f} €"


def dados_vazios():
    """Estrutura de dados em branco, usada quando a pessoa prefere
    preencher tudo manualmente em vez de tirar uma foto."""
    return {
        "NomeCliente": "",
        "MoradaCliente": "",
        "Pagamento": "",
        "Itens": [
            {"Designação": "", "Unidade": "Vg.", "Quantidade": 1.0, "Preço Unitário (€)": 0.0}
        ],
    }


def normalizar_dados(dados):
    """Garante que os dados devolvidos pela IA têm sempre a estrutura e os
    tipos esperados, mesmo que a IA se engane ou omita campos."""
    if not isinstance(dados, dict):
        dados = {}

    itens_normalizados = []
    for item in dados.get("Itens", []) or []:
        if not isinstance(item, dict):
            continue
        designacao = str(item.get("Designação") or "").strip()
        unidade = str(item.get("Unidade") or "Vg.").strip() or "Vg."
        itens_normalizados.append(
            {
                "Designação": designacao,
                "Unidade": unidade,
                "Quantidade": parse_numero(item.get("Quantidade", 1), default=1.0),
                "Preço Unitário (€)": parse_numero(item.get("Preço Unitário (€)", 0)),
            }
        )

    if not itens_normalizados:
        itens_normalizados = [
            {"Designação": "", "Unidade": "Vg.", "Quantidade": 1.0, "Preço Unitário (€)": 0.0}
        ]

    return {
        "NomeCliente": str(dados.get("NomeCliente") or "").strip(),
        "MoradaCliente": str(dados.get("MoradaCliente") or "").strip(),
        "Pagamento": str(dados.get("Pagamento") or "").strip(),
        "Itens": itens_normalizados,
    }


def preparar_imagem_para_ia(imagem):
    """Reduz o tamanho da fotografia antes de a enviar para a IA. As fotos
    de telemóveis atuais são muito grandes e isso torna a análise mais
    lenta sem ganho de qualidade na leitura do texto."""
    imagem = imagem.convert("RGB")
    largura_max = 1600
    if imagem.width > largura_max:
        proporcao = largura_max / imagem.width
        nova_altura = int(imagem.height * proporcao)
        imagem = imagem.resize((largura_max, nova_altura))
    return imagem


# =========================================================================
# INTELIGÊNCIA ARTIFICIAL
# =========================================================================

def obter_cliente_ia():
    """Cria (uma única vez por sessão) o cliente do Gemini."""
    if "cliente_ia" not in st.session_state:
        if not API_KEY:
            raise RuntimeError("A chave da API de IA não está configurada.")
        st.session_state.cliente_ia = genai.Client(api_key=API_KEY)
    return st.session_state.cliente_ia


def extrair_dados_da_imagem(imagem):
    """Envia a imagem para o Gemini e devolve os dados já normalizados."""
    imagem_preparada = preparar_imagem_para_ia(imagem)
    cliente = obter_cliente_ia()
    prompt = """
    Lê atentamente as anotações manuscritas desta imagem, referentes a um
    orçamento de obra ou serviço.

    Extrai o nome do cliente, a morada do cliente (se existirem), as
    condições de pagamento e a lista de trabalhos/materiais.

    Devolve APENAS um objeto JSON válido, sem texto adicional, sem
    comentários e sem marcações ```json, com esta estrutura exata:
    {
      "NomeCliente": "nome do cliente ou string vazia",
      "MoradaCliente": "morada do cliente ou string vazia",
      "Pagamento": "condições de pagamento ou string vazia",
      "NomeEmpresa": "nome da empresa / pessoa ou string vazia",
      "Contato": "telefone ou telemóvel ou string vazia",
      "Email": "email ou string vazia",
      "Itens": [
        {
          "Designação": "descrição do trabalho ou material",
          "Unidade": "Vg.",
          "Quantidade": 1,
          "Preço Unitário (€)": 0.00
        }
      ]
    }

    Regras:
    - "Quantidade" e "Preço Unitário (€)" devem ser números, não texto.
    - Se não existir uma quantidade explícita, usa 1.
    - "Unidade" pode ser "Vg." (verba), "un.", "m2", "m", "h" ou outra
      unidade indicada nas anotações.
    """
    resposta = cliente.models.generate_content(
        model=NOME_MODELO_IA, contents=[prompt, imagem_preparada]
    )

    texto = resposta.text.strip()
    texto = texto.replace("```json", "").replace("```", "").strip()

    try:
        dados = json.loads(texto)
    except json.JSONDecodeError:
        inicio, fim = texto.find("{"), texto.rfind("}")
        if inicio == -1 or fim == -1:
            raise
        dados = json.loads(texto[inicio : fim + 1])

    return normalizar_dados(dados)


# =========================================================================
# GERAÇÃO DO DOCUMENTO (Word → PDF)
# =========================================================================

def substituir_na_paragrafo(paragrafo, mapa_substituicoes):
    """Substitui as etiquetas {{...}} no texto de um parágrafo, preservando
    a formatação (negrito, tamanho, tipo de letra) do texto original."""
    texto_completo = paragrafo.text
    if not any(chave in texto_completo for chave in mapa_substituicoes):
        return
    for chave, valor in mapa_substituicoes.items():
        texto_completo = texto_completo.replace(chave, valor)

    if paragrafo.runs:
        paragrafo.runs[0].text = texto_completo
        for run in paragrafo.runs[1:]:
            run.text = ""
    else:
        paragrafo.text = texto_completo


def obter_pasta_trabalho():
    """Cada sessão (cada pessoa a usar a app) tem a sua própria pasta
    temporária, para que dois orçamentos gerados ao mesmo tempo por
    pessoas diferentes nunca se misturem."""
    if "pasta_trabalho" not in st.session_state:
        st.session_state.pasta_trabalho = tempfile.mkdtemp(prefix="orc_")
    return st.session_state.pasta_trabalho


def gerar_documento(nome_cliente, morada_cliente, dataframe_tabela, pagamento, nome_empresa, contato, email):
    """Recebe os dados confirmados, preenche o Word e converte para PDF
    usando o LibreOffice."""
    pasta_trabalho = obter_pasta_trabalho()

    doc = Document(TEMPLATE_PATH)
    data_hoje = datetime.today().strftime("%d/%m/%Y")

    substituicoes = {
        "{{NOME_CLIENTE}}": nome_cliente,
        "{{MORADA}}": morada_cliente,
        "{{PAGAMENTO}}": pagamento,
        "{{DATA}}": data_hoje,
        "{{NOME_EMPRESA}}": nome_empresa,
        "{{CONTATO}}": contato,
        "{{EMAIL}}": email,
    }
    for paragrafo in doc.paragraphs:
        substituir_na_paragrafo(paragrafo, substituicoes)

    tabela_word = doc.tables[0]
    total_orcamento = 0.0

    # Guardar as estruturas XML dos dois moldes (linha de item e linha de total)
    molde_item = deepcopy(tabela_word.rows[2]._tr)
    molde_total = deepcopy(tabela_word.rows[3]._tr)

    # Limpar a tabela original, ficando só com o cabeçalho
    tabela_word._tbl.remove(tabela_word.rows[3]._tr)
    tabela_word._tbl.remove(tabela_word.rows[2]._tr)
    tabela_word._tbl.remove(tabela_word.rows[1]._tr)

    numero_item = 0
    for _, linha in dataframe_tabela.iterrows():
        designacao = str(linha.get("Designação") or "").strip()
        if not designacao:
            continue  # ignora linhas em branco deixadas na tabela

        numero_item += 1
        nova_linha_tr = deepcopy(molde_item)
        tabela_word._tbl.append(nova_linha_tr)
        celulas = tabela_word.rows[-1].cells
        for celula in celulas:
            celula.text = ""

        quantidade = parse_numero(linha.get("Quantidade", 0), default=0.0)
        preco_unitario = parse_numero(linha.get("Preço Unitário (€)", 0), default=0.0)
        preco_total_linha = quantidade * preco_unitario
        total_orcamento += preco_total_linha

        celulas[0].text = str(numero_item)
        celulas[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas[1].text = designacao
        celulas[2].text = str(linha.get("Unidade") or "Vg.")
        celulas[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas[3].text = formatar_numero(quantidade)
        celulas[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas[4].text = formatar_numero(preco_unitario)
        celulas[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        celulas[5].text = formatar_numero(preco_total_linha)
        celulas[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linha final de total
    nova_linha_total_tr = deepcopy(molde_total)
    tabela_word._tbl.append(nova_linha_total_tr)
    linha_total = tabela_word.rows[-1].cells
    linha_total[4].text = "TOTAL:"
    linha_total[5].text = f"{total_orcamento:.2f} €"
    if linha_total[4].paragraphs[0].runs:
        linha_total[4].paragraphs[0].runs[0].bold = True
    if linha_total[5].paragraphs[0].runs:
        linha_total[5].paragraphs[0].runs[0].bold = True

    docx_path = os.path.join(pasta_trabalho, "orcamento.docx")
    doc.save(docx_path)

    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", pasta_trabalho, docx_path],
        check=True,
        timeout=90,
    )
    pdf_path = os.path.join(pasta_trabalho, "orcamento.pdf")
    return pdf_path, data_hoje


# =========================================================================
# ECRÃS
# =========================================================================

def ecra_login():
    st.markdown('<div class="login-card">', unsafe_allow_html=True)
    st.markdown("### Acesso Reservado")
    st.markdown("Introduza a password para continuar.")
    senha = st.text_input(
        "Password", type="password", label_visibility="collapsed", placeholder="Password"
    )
    entrar = st.button("Entrar", use_container_width=True, type="primary")
    if entrar:
        if PASSWORD_SISTEMA and senha == PASSWORD_SISTEMA:
            st.session_state.autenticado = True
            st.rerun()
        else:
            st.error("Password incorreta. Tente novamente.")
    st.markdown("</div>", unsafe_allow_html=True)


def passo_1_foto():
    st.markdown("## Passo 1 - Foto do Orçamento")
    st.markdown(
        '<div class="dica-caixa">Tire a foto num local bem iluminado, '
        "com o papel esticado e sem sombras por cima do texto.</div>",
        unsafe_allow_html=True,
    )

    aba_camera, aba_ficheiro = st.tabs(["Tirar Foto", "Escolher Ficheiro"])
    with aba_camera:
        foto_camera = st.camera_input("Tirar foto", label_visibility="collapsed")
    with aba_ficheiro:
        foto_ficheiro = st.file_uploader(
            "Escolher imagem", type=["jpg", "jpeg", "png"], label_visibility="collapsed"
        )

    arquivo_imagem = foto_ficheiro
    if foto_camera is not None:
        arquivo_imagem = foto_camera

    if arquivo_imagem is not None:
        imagem = Image.open(arquivo_imagem)
        st.image(imagem, caption="Pré-visualização", use_container_width=True)

        if st.button("Confirmar e Analisar", use_container_width=True, type="primary"):
            with st.spinner("A analisar o orçamento... alguns segundos."):
                try:
                    dados = extrair_dados_da_imagem(imagem)
                    st.session_state.dados_extraidos = dados
                    st.session_state.passo = 2
                    st.rerun()
                except json.JSONDecodeError:
                    st.error(
                        "Não conseguimos interpretar os dados da foto. "
                        "Tente novamente com mais luz e o papel bem esticado."
                    )
                except Exception as e:
                    st.error("Ocorreu um erro ao analisar a foto. Tente novamente.")
                    with st.expander("Detalhes técnicos"):
                        st.code(str(e))

    st.markdown("---")
    if st.button("Prefiro preencher os dados manualmente"):
        st.session_state.dados_extraidos = dados_vazios()
        st.session_state.passo = 2
        st.rerun()


def passo_2_confirmar():
    if "dados_extraidos" not in st.session_state:
        st.session_state.passo = 1
        st.rerun()

    st.markdown("## Passo 2 - Confirme os Dados")
    dados = st.session_state.dados_extraidos

    st.markdown("#### Os seus dados")
    nome_empresa = st.text_input("Nome da Empresa / Pessoa", value=dados.get("NomeEmpresa", ""))
    contato = st.text_input("Contato", value=dados.get("Contato", ""))
    email = st.text_input("Email", value=dados.get("Email", ""))

    st.markdown("#### Dados do Cliente")
    nome_cliente = st.text_input("Nome do Cliente", value=dados.get("NomeCliente", ""))
    morada_cliente = st.text_area("Morada do Cliente", value=dados.get("MoradaCliente", ""))

    st.markdown("#### Trabalhos e Materiais")
    st.caption("Reveja as descrições, quantidades e preços. Pode adicionar ou apagar linhas.")

    df = pd.DataFrame(dados.get("Itens", []))
    for coluna in COLUNAS_TABELA:
        if coluna not in df.columns:
            df[coluna] = "" if coluna in ("Designação", "Unidade") else 0.0
    df = df[COLUNAS_TABELA]

    tabela_editada = st.data_editor(
        df,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        key="editor_tabela",
        column_config={
            "Designação": st.column_config.TextColumn("Descrição", width="large"),
            "Unidade": st.column_config.TextColumn("Unid.", width="small"),
            "Quantidade": st.column_config.NumberColumn("Qtd.", min_value=0.0, step=0.5, format="%.2f"),
            "Preço Unitário (€)": st.column_config.NumberColumn(
                "Preço Unit. (€)", min_value=0.0, step=0.5, format="%.2f"
            ),
        },
    )

    quantidades = tabela_editada["Quantidade"].apply(parse_numero)
    precos = tabela_editada["Preço Unitário (€)"].apply(parse_numero)
    total = (quantidades * precos).sum()
    st.metric("Total do Orçamento", formatar_euro(total))

    pagamento = st.text_input("Condições de Pagamento", value=dados.get("Pagamento", ""))

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Voltar", use_container_width=True):
            st.session_state.passo = 1
            st.rerun()
    with col2:
        gerar = st.button("Gerar PDF", use_container_width=True, type="primary")

    if gerar:
        if not morada_cliente.strip():
            st.warning("Por favor, preencha a morada do cliente.")
        elif tabela_editada["Designação"].apply(lambda x: str(x or "").strip()).eq("").all():
            st.warning("Adicione pelo menos um trabalho ou material ao orçamento.")
        else:
            with st.spinner("A gerar o documento Word e PDF..."):
                try:
                    pdf_path, data_doc = gerar_documento(
                        nome_cliente, morada_cliente, tabela_editada, pagamento, nome_empresa, contato, email
                    )
                    st.session_state.pdf_path = pdf_path
                    st.session_state.data_doc = data_doc
                    st.session_state.nome_cliente_final = nome_cliente
                    st.session_state.passo = 3
                    st.rerun()
                except subprocess.CalledProcessError:
                    st.error("Não foi possível converter o documento para PDF. Tente novamente.")
                except Exception as e:
                    st.error("Ocorreu um erro ao gerar o documento.")
                    with st.expander("Detalhes técnicos"):
                        st.code(str(e))


def passo_3_download():
    if "pdf_path" not in st.session_state:
        st.session_state.passo = 1
        st.rerun()

    st.markdown("##Orçamento Pronto")
    st.success("O PDF foi gerado com sucesso.")

    nome_base = (st.session_state.get("nome_cliente_final") or "cliente").strip()
    nome_base = re.sub(r"\s+", "_", nome_base) or "cliente"
    nome_ficheiro = f"Orcamento_{nome_base}_{st.session_state.data_doc.replace('/', '-')}.pdf"

    with open(st.session_state.pdf_path, "rb") as ficheiro_pdf:
        st.download_button(
            label="Descarregar PDF",
            data=ficheiro_pdf,
            file_name=nome_ficheiro,
            mime="application/pdf",
            use_container_width=True,
            type="primary",
        )

    st.markdown("---")
    if st.button("Criar Novo Orçamento", use_container_width=True):
        reiniciar_sessao()
        st.rerun()


def reiniciar_sessao():
    for chave in ["dados_extraidos", "pdf_path", "data_doc", "nome_cliente_final"]:
        st.session_state.pop(chave, None)
    st.session_state.passo = 1


# =========================================================================
# ARRANQUE DA APLICAÇÃO
# =========================================================================

def main():
    st.set_page_config(page_title="Orçamentos", page_icon="🧾", layout="centered")
    aplicar_estilo()

    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False

    if not st.session_state.autenticado:
        st.markdown("<h1>Orçamentos</h1>", unsafe_allow_html=True)
        ecra_login()
        return

    if "passo" not in st.session_state:
        st.session_state.passo = 1

    st.markdown("<h1>Orçamentos</h1>", unsafe_allow_html=True)
    indicador_passos(st.session_state.passo)

    if st.session_state.passo == 1:
        passo_1_foto()
    elif st.session_state.passo == 2:
        passo_2_confirmar()
    elif st.session_state.passo == 3:
        passo_3_download()


if __name__ == "__main__":
    main()